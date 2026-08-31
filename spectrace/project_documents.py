"""Safe text extraction and candidate-scope contracts for New Project Beta."""

from __future__ import annotations

import re
from zipfile import BadZipFile
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path, PurePath
from typing import Protocol

from docx import Document
from pydantic import Field
from pypdf import PdfReader
from pypdf.errors import PdfReadError

from spectrace.models import StrictModel
from spectrace.workflow import StructuredWorkflowStep


SUPPORTED_SUFFIXES = frozenset({".pdf", ".docx", ".md", ".txt"})
MAX_DOCUMENT_BYTES = 8 * 1024 * 1024


class DocumentExtractionError(ValueError):
    """A user-correctable document validation or extraction failure."""


class CandidateScopeItem(StrictModel):
    text: str = Field(min_length=1, max_length=2000)
    source_filename: str = Field(min_length=1, max_length=255)
    source_location: str = Field(min_length=1, max_length=120)
    supporting_quote: str = Field(min_length=1, max_length=2000)
    category: str = Field(pattern=r"^(APPROVED_REQUIREMENT|CONSTRAINT|EXCLUSION|ASSUMPTION|UNRESOLVED_QUESTION|DECISION)$")
    confidence: float = Field(ge=0, le=1)
    uncertainty: str | None = Field(default=None, max_length=500)


class CandidateScopeExtraction(StrictModel):
    project_name: str = Field(min_length=1, max_length=160)
    items: tuple[CandidateScopeItem, ...]


class CandidateWorkflow(StrictModel):
    """A deterministic candidate found only in an explicit approved-workflow section."""

    steps: tuple[StructuredWorkflowStep, ...]
    exception_branches: tuple[str, ...] = ()
    source_filename: str = Field(min_length=1, max_length=255)
    source_location: str = Field(min_length=1, max_length=120)


@dataclass(frozen=True)
class ExtractedSection:
    source_filename: str
    source_location: str
    text: str


@dataclass(frozen=True)
class ExtractedDocument:
    filename: str
    sections: tuple[ExtractedSection, ...]

    @property
    def text(self) -> str:
        return "\n\n".join(section.text for section in self.sections)


class CandidateScopeClient(Protocol):
    def generate_candidate_scope(
        self, documents: tuple[ExtractedDocument, ...], project_name: str
    ) -> CandidateScopeExtraction: ...


def safe_filename(value: str) -> str:
    """Drop path components and restrict names used in local display/storage."""

    name = PurePath(value.replace("\\", "/")).name
    safe = re.sub(r"[^A-Za-z0-9._ -]", "_", name).strip(" .")
    if not safe:
        raise DocumentExtractionError("The selected file does not have a usable filename.")
    return safe[:255]


def validate_document(filename: str, content: bytes) -> tuple[str, str]:
    name = safe_filename(filename)
    suffix = Path(name).suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        raise DocumentExtractionError("Use a PDF, DOCX, Markdown or text document.")
    if not content:
        raise DocumentExtractionError("The selected file is empty.")
    if len(content) > MAX_DOCUMENT_BYTES:
        raise DocumentExtractionError("The selected file exceeds the allowed 8 MB size.")
    return name, suffix


def _text_sections(name: str, text: str) -> tuple[ExtractedSection, ...]:
    text = text.replace("\r\n", "\n").replace("\r", "\n").strip()
    if not text:
        raise DocumentExtractionError("This document contains no extractable text.")
    lines = text.splitlines()
    sections: list[ExtractedSection] = []
    start = 1
    buffer: list[str] = []
    for line_number, line in enumerate(lines, start=1):
        if line.startswith("#") and buffer:
            value = "\n".join(buffer).strip()
            if value:
                sections.append(ExtractedSection(name, f"lines {start}–{line_number - 1}", value))
            start, buffer = line_number, [line]
        else:
            buffer.append(line)
    value = "\n".join(buffer).strip()
    if value:
        sections.append(ExtractedSection(name, f"lines {start}–{len(lines)}", value))
    return tuple(sections)


def extract_document(filename: str, content: bytes) -> ExtractedDocument:
    """Extract text in memory; uploaded bytes are never written to Git paths."""

    name, suffix = validate_document(filename, content)
    try:
        if suffix in {".txt", ".md"}:
            sections = _text_sections(name, content.decode("utf-8-sig"))
        elif suffix == ".docx":
            document = Document(BytesIO(content))
            sections = tuple(
                ExtractedSection(name, f"paragraph {index}", paragraph.text.strip())
                for index, paragraph in enumerate(document.paragraphs, start=1)
                if paragraph.text.strip()
            )
        else:
            reader = PdfReader(BytesIO(content))
            sections = tuple(
                ExtractedSection(name, f"page {index}", text.strip())
                for index, page in enumerate(reader.pages, start=1)
                if (text := (page.extract_text() or "")).strip()
            )
    except (UnicodeDecodeError, OSError, ValueError, BadZipFile, PdfReadError) as exc:
        raise DocumentExtractionError("The selected document could not be read as text.") from exc
    if not sections:
        message = (
            "This PDF contains no extractable text. Scanned-document OCR is not supported yet."
            if suffix == ".pdf"
            else "This document contains no extractable text."
        )
        raise DocumentExtractionError(message)
    return ExtractedDocument(filename=name, sections=sections)


def render_candidate_prompt(
    documents: tuple[ExtractedDocument, ...], project_name: str
) -> str:
    """Create a strict extraction prompt containing source locations and text only."""

    blocks = []
    for document in documents:
        for section in document.sections:
            blocks.append(
                f"SOURCE: {section.source_filename}\nLOCATION: {section.source_location}\nTEXT:\n{section.text}"
            )
    return (
        "Extract a candidate project scope for human review. Preserve an exact supporting quote "
        "and its source filename/location for every item. Never promote assumptions or unresolved "
        "questions into approved requirements. Return only the required structured object.\n\n"
        f"PROJECT: {project_name}\n\n" + "\n\n---\n\n".join(blocks)
    )


def candidate_from_structured_markdown(
    document: ExtractedDocument, project_name: str
) -> CandidateScopeExtraction:
    """Load the bundled synthetic example without a provider call."""

    category_by_heading = {
        "approved requirements": "APPROVED_REQUIREMENT",
        "constraints": "CONSTRAINT",
        "exclusions": "EXCLUSION",
        "assumptions": "ASSUMPTION",
        "unresolved questions": "UNRESOLVED_QUESTION",
        "decisions": "DECISION",
    }
    items: list[CandidateScopeItem] = []
    for section in document.sections:
        lines = section.text.splitlines()
        if not lines:
            continue
        heading = lines[0].lstrip("# ").strip().lower()
        category = category_by_heading.get(heading)
        if not category:
            continue
        for line in lines[1:]:
            quote = re.sub(r"^[-*]\s+", "", line).strip()
            if not quote:
                continue
            items.append(
                CandidateScopeItem(
                    text=quote,
                    source_filename=document.filename,
                    source_location=section.source_location,
                    supporting_quote=quote,
                    category=category,
                    confidence=1.0,
                )
            )
    if not items:
        raise DocumentExtractionError("The synthetic example did not contain structured scope sections.")
    return CandidateScopeExtraction(project_name=project_name, items=tuple(items))


_APPROVED_WORKFLOW_HEADING = re.compile(
    r"^(?:#{1,6}\s*)?(?:\d+[.)]\s*)?(?:original\s+)?approved(?:\s+end-to-end)?(?:\s+business)?\s+workflow\s*:?[ \t]*$",
    re.IGNORECASE,
)
_NUMBERED_WORKFLOW_STEP = re.compile(
    r"^\s*\d+[.)]\s*(?P<actor>[A-Za-z][A-Za-z /&-]{0,118})\s*:\s*(?P<action>\S.*)$"
)
_EXCEPTION_BRANCH = re.compile(
    r"^\s*(?:[-*]\s*)?(?:exception|otherwise|if\s+.+?|when\s+.+?)\s*:\s*(?P<action>\S.*)$",
    re.IGNORECASE,
)


def extract_explicit_approved_workflow(
    documents: tuple[ExtractedDocument, ...],
) -> CandidateWorkflow | None:
    """Parse an explicit approved workflow; never infer a process from ordinary prose."""

    for document in documents:
        lines_with_location = [
            (line, section.source_location)
            for section in document.sections
            for line in section.text.splitlines()
        ]
        heading_index = next(
            (
                index for index, (line, _) in enumerate(lines_with_location)
                if _APPROVED_WORKFLOW_HEADING.match(line.strip())
            ),
            None,
        )
        if heading_index is None:
            continue
        steps: list[StructuredWorkflowStep] = []
        exception_branches: list[str] = []
        in_exceptions = False
        for line, _ in lines_with_location[heading_index + 1 :]:
                normalized = line.strip()
                if re.match(r"^(?:#{1,6}\s*)?exception branches?\s*$", normalized, re.IGNORECASE):
                    in_exceptions = True
                    continue
                if in_exceptions:
                    if re.match(r"^(?:#{1,6}\s*)?(?:acceptance|constraints?|exclusions?|assumptions?|decisions?|unresolved)\b", normalized, re.IGNORECASE):
                        break
                    if ":" in normalized:
                        exception_branches.append(normalized)
                    continue
                if line.lstrip().startswith("#") and steps:
                    break
                match = _NUMBERED_WORKFLOW_STEP.match(line)
                if match:
                    action = match.group("action").strip()
                    branch = None
                    bracketed = re.search(r"\s+\[(.+)]\s*$", action)
                    if bracketed:
                        branch = bracketed.group(1).strip()
                        action = action[: bracketed.start()].strip()
                    steps.append(
                        StructuredWorkflowStep(
                            actor=match.group("actor").strip(), action=action, branch=branch
                        )
                    )
                    continue
                branch_match = _EXCEPTION_BRANCH.match(line)
                if branch_match and steps:
                    prior = steps[-1]
                    branch_text = branch_match.group("action").strip()
                    steps[-1] = prior.model_copy(
                        update={"branch": f"{prior.branch}; {branch_text}" if prior.branch else branch_text}
                    )
                # Introductory prose between the heading and first numbered step is ignored.
        if steps:
            return CandidateWorkflow(
                steps=tuple(steps),
                exception_branches=tuple(exception_branches),
                source_filename=document.filename,
                source_location=lines_with_location[heading_index][1],
            )
    return None


def friendly_validation_message(exc: Exception) -> str:
    """Map validation failures to concise UI copy without model internals or URLs."""

    text = str(exc).lower()
    if "project_name" in text or "project name" in text:
        return "Enter a project name."
    if "approved_requirements" in text or "approved requirement" in text:
        return "Add at least one approved requirement."
    if "date" in text:
        return "Review the decision date format: use YYYY-MM-DD."
    if "workflow" in text and ("actor" in text or "action" in text):
        return "This workflow step needs both an actor and an action."
    if isinstance(exc, DocumentExtractionError):
        return str(exc)
    return "Review the highlighted project information and try again."


def friendly_validation_messages(exc: Exception, *, limit: int = 3) -> tuple[str, ...]:
    """Return at most three deduplicated, business-facing validation messages."""

    raw = str(exc).lower()
    messages: list[str] = []
    checks = (
        (("project_name", "project name"), "Enter a project name."),
        (("approved_requirements", "approved requirement"), "Add at least one approved requirement."),
        (("date", "date_from"), "Review the decision date format: use YYYY-MM-DD."),
        (("workflow", "actor", "action"), "This workflow step needs both an actor and an action."),
    )
    for terms, message in checks:
        if any(term in raw for term in terms) and message not in messages:
            messages.append(message)
    if not messages:
        messages.append(friendly_validation_message(exc))
    return tuple(messages[:limit])
