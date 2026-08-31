from __future__ import annotations

from datetime import date
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from pypdf import PdfReader, PdfWriter
from pypdf.generic import DictionaryObject, NameObject, StreamObject

from spectrace.beta_project import materialize_beta_project
from spectrace.project_documents import (
    MAX_DOCUMENT_BYTES,
    DocumentExtractionError,
    candidate_from_structured_markdown,
    extract_explicit_approved_workflow,
    extract_document,
    friendly_validation_message,
    friendly_validation_messages,
    render_candidate_prompt,
)
from spectrace.scope_anchor import build_scope_anchor
from spectrace.workflow import (
    StructuredDecision,
    StructuredProjectInput,
    build_structured_scope_anchor,
    generate_workflow_draft,
    verify_workflow_draft,
)
from spectrace.ledger import LedgerStore


ROOT = Path(__file__).resolve().parents[1]


def _docx_bytes() -> bytes:
    document = Document()
    document.add_heading("Approved requirements", level=1)
    document.add_paragraph("A fictional student can browse rooms.")
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def _pdf_bytes(text: str | None = None) -> bytes:
    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)
    if text:
        font = DictionaryObject(
            {
                NameObject("/Type"): NameObject("/Font"),
                NameObject("/Subtype"): NameObject("/Type1"),
                NameObject("/BaseFont"): NameObject("/Helvetica"),
            }
        )
        font_ref = writer._add_object(font)
        page[NameObject("/Resources")] = DictionaryObject(
            {NameObject("/Font"): DictionaryObject({NameObject("/F1"): font_ref})}
        )
        content = StreamObject()
        content.set_data(f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode("ascii"))
        page[NameObject("/Contents")] = writer._add_object(content)
    stream = BytesIO()
    writer.write(stream)
    return stream.getvalue()


def test_txt_markdown_and_multiple_document_extraction_preserve_sources() -> None:
    markdown = extract_document("scope.md", b"# Scope\n- Student can browse rooms.\n# Limits\n- One campus.")
    text = extract_document("notes.txt", b"Decision notes\nApproved for the fictional beta.")
    assert markdown.filename == "scope.md"
    assert markdown.sections[0].source_location == "lines 1\u20132"
    assert text.sections[0].source_filename == "notes.txt"
    prompt = render_candidate_prompt((markdown, text), "CampusFlow")
    assert "SOURCE: scope.md" in prompt and "SOURCE: notes.txt" in prompt
    assert "Student can browse rooms" in prompt


def test_docx_and_text_pdf_extraction_preserve_paragraph_and_page() -> None:
    docx = extract_document("requirements.docx", _docx_bytes())
    pdf = extract_document("proposal.pdf", _pdf_bytes("Fictional approved room booking"))
    assert any(section.source_location == "paragraph 2" for section in docx.sections)
    assert pdf.sections[0].source_location == "page 1"
    assert "Fictional approved room booking" in pdf.text


def test_no_text_pdf_type_and_size_fail_with_friendly_messages() -> None:
    with pytest.raises(DocumentExtractionError, match="OCR is not supported"):
        extract_document("scan.pdf", _pdf_bytes())
    with pytest.raises(DocumentExtractionError, match="PDF, DOCX"):
        extract_document("scope.exe", b"not executable")
    with pytest.raises(DocumentExtractionError, match="8 MB"):
        extract_document("large.txt", b"x" * (MAX_DOCUMENT_BYTES + 1))


def test_synthetic_candidate_requires_human_approval_and_materializes_only_locally(tmp_path: Path) -> None:
    sample = ROOT / "assets" / "CampusFlow_Synthetic_Project.md"
    document = extract_document(sample.name, sample.read_bytes())
    candidate = candidate_from_structured_markdown(document, "CampusFlow")
    requirements = tuple(item.text for item in candidate.items if item.category == "APPROVED_REQUIREMENT")
    project = StructuredProjectInput(
        project_name="CampusFlow",
        approved_requirements=requirements,
        decisions=(StructuredDecision(effective_date=date(2026, 8, 31), text="Approve fictional scope."),),
    )
    with pytest.raises(ValueError, match="explicit human approval"):
        build_structured_scope_anchor(project, human_approved=False)
    approved = build_structured_scope_anchor(project, human_approved=True)
    assert approved.items
    pack = materialize_beta_project(project, tmp_path, "unit-test")
    parsed = build_scope_anchor(pack)
    assert parsed.project_id == "campusflow"
    assert pack.is_relative_to(tmp_path)
    assert not (ROOT / "CampusFlow_Synthetic_Project.md").exists()


def test_friendly_validation_never_leaks_pydantic_urls_or_internal_fields() -> None:
    message = friendly_validation_message(ValueError("string_too_short project_name https://errors.pydantic.dev/2.11/v/string_too_short"))
    assert message == "Enter a project name."
    assert "pydantic" not in message.lower()
    assert "string_too_short" not in message


def test_explicit_approved_workflow_parser_preserves_actor_action_and_exception() -> None:
    document = extract_document(
        "HarborBasket.md",
        b"# Approved end-to-end business workflow\n"
        b"1. Member: Submit a pickup request\n"
        b"2. System: Validate available inventory [If unavailable, show an exception]\n"
        b"3. Coordinator: Confirm the collection window\n"
        b"# Assumptions\n- None\n",
    )
    candidate = extract_explicit_approved_workflow((document,))
    assert candidate is not None
    assert [(step.actor, step.action) for step in candidate.steps] == [
        ("Member", "Submit a pickup request"),
        ("System", "Validate available inventory"),
        ("Coordinator", "Confirm the collection window"),
    ]
    assert candidate.steps[1].branch == "If unavailable, show an exception"


def test_workflow_parser_refuses_to_infer_from_unstructured_prose() -> None:
    document = extract_document("notes.txt", b"Members submit requests and the system validates them.")
    assert extract_explicit_approved_workflow((document,)) is None


def test_approved_explicit_workflow_materializes_and_verifies(tmp_path: Path) -> None:
    document = extract_document(
        "HarborBasket.md",
        b"# Approved business workflow\n1. Member: Submit a pickup request\n2. System: Validate inventory\n3. Coordinator: Confirm pickup\n",
    )
    candidate = extract_explicit_approved_workflow((document,))
    assert candidate is not None
    project = StructuredProjectInput(
        project_name="HarborBasket",
        approved_requirements=("Members can request a pickup.",),
        workflow_steps=candidate.steps,
    )
    pack = materialize_beta_project(project, tmp_path, "workflow-test")
    anchor = build_scope_anchor(pack)
    with LedgerStore() as ledger:
        ledger.seed_anchor(anchor, pack, approved_through="DEC-001")
        draft = generate_workflow_draft(anchor, pack, ledger, evidence_cutoff="DEC-001")
        verification = verify_workflow_draft(draft, anchor, pack, ledger.snapshot(anchor.project_id))
    assert verification.passed
    assert [node.actor_id for node in draft.nodes[1:-1]] == [
        "ACTOR-MEMBER", "ACTOR-SYSTEM", "ACTOR-COORDINATOR"
    ]


def test_validation_messages_are_bounded_and_sanitized() -> None:
    messages = friendly_validation_messages(
        ValueError("project_name approved_requirements invalid date workflow actor https://errors.pydantic.dev")
    )
    assert len(messages) == 3
    assert all("pydantic" not in message.lower() and "http" not in message.lower() for message in messages)
